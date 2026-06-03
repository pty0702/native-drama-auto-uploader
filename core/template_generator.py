"""
成本配置模板生成器

流程：
1. 用 python-docx 在 视频.docx 模板中填入微短剧名称、集数和总时长
2. 通过 Word COM → PDF → PyMuPDF 转 jpg
3. 将填好字的 docx 图片、带章模板图一起发送 images/edits API
4. AI 按提示词融合生成一张真实打印纸放在桌面上的自然照片，作为最终模版.jpg
"""
from __future__ import annotations

import base64
import math
import mimetypes
import os
import queue
import random
import shutil
import tempfile
import threading
from io import BytesIO
from pathlib import Path

import requests
from docx import Document
from PIL import Image, ImageDraw, ImageFilter, ImageFont, ImageOps

from native_drama_uploader.settings import IMAGE_API_BASE_URL, IMAGE_MODEL, SUCAI_DIR

# 绕过系统代理
_session = requests.Session()
_session.trust_env = False


def _normalize_api_base_url(api_base_url):
    api_base_url = (api_base_url or IMAGE_API_BASE_URL).rstrip("/")
    if not api_base_url:
        raise RuntimeError("未配置图片 API 地址")
    if not api_base_url.endswith("/v1"):
        api_base_url = f"{api_base_url}/v1"
    return api_base_url


def duration_minutes_for_report(total_duration_sec: float) -> int:
    """成本配置表按分钟向上取整，不满 1 分钟按 1 分钟填写。"""
    if total_duration_sec <= 0:
        return 1
    return max(1, math.ceil(total_duration_sec / 60))


def generate_template_image(
    template_path: str | Path,
    output_dir: str,
    sub_name: str,
    episode_count: int,
    total_duration_sec: float,
    stamp_image: str | Path | None = None,
    api_key: str | None = None,
    api_base_url: str | None = None,
    image_model: str | None = None,
    max_attempts: int = 3,
    log_cb=None,
) -> str:
    """
    生成成本配置比例情况报告的自然照片。

    template_path: 视频.docx 模板路径
    stamp_image: 模板.jpg / 图2.jpg（带承诺文字和印章的参考图）路径
    """
    def log(msg):
        if log_cb:
            log_cb(msg)

    os.makedirs(output_dir, exist_ok=True)
    template_path = str(template_path)
    stamp_image_path = _resolve_stamp_image(template_path, stamp_image)
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"成本模板 docx 不存在: {template_path}")
    if not os.path.exists(stamp_image_path):
        raise FileNotFoundError(f"印章底图不存在: {stamp_image_path}")

    # ---- 步骤 0: 从 docx 提取报审机构 ----
    company_name = _extract_company_from_docx(template_path)
    if company_name:
        log(f"从 docx 提取报审机构: {company_name}")
    else:
        log("未从 docx 提取到报审机构，AI 将自行判断")

    # ---- 步骤 1: 填写 docx ----
    log("步骤 1: 填写成本配置模板...")
    total_min = duration_minutes_for_report(total_duration_sec)
    filled_docx = _fill_docx(template_path, sub_name, episode_count, total_min)

    # ---- 步骤 2: docx → jpg ----
    log("步骤 2: 转换 docx 为图片...")
    docx_img = _docx_to_image(filled_docx, log_cb=log)

    docx_image_path = os.path.join(output_dir, "_docx_filled_input.jpg")
    docx_img.save(docx_image_path, "JPEG", quality=95)
    log(f"  已保存填字后的 docx 图片: {docx_image_path}")

    # 旧方案已停用：不再用 docx 上半部分 + 模板下半部分做本地拼接。
    # merged_img = _merge_images(docx_img, stamp_image_path)
    # merged_path = os.path.join(output_dir, "_merged_input.jpg")
    # merged_img.save(merged_path, "JPEG", quality=95)

    # 旧兜底已停用：不再直接擦写 模板.jpg 的表格区域。
    # stable_input = _fill_stamp_template(stamp_image_path, output_dir, sub_name, episode_count, total_min, log)

    # ---- 步骤 3: AI 融合 docx 和印章模板 ----
    active_model = image_model or IMAGE_MODEL
    log(f"步骤 3: 调用 AI({active_model}) 融合 docx 和印章模板...")
    ai_img = _ai_fuse_template_photo(
        docx_image_path=docx_image_path,
        stamp_image_path=stamp_image_path,
        sub_name=sub_name,
        episode_count=episode_count,
        total_min=total_min,
        company_name=company_name,
        api_key=api_key,
        api_base_url=api_base_url,
        image_model=active_model,
        log_cb=log,
    )
    out_path = os.path.join(output_dir, "模版.jpg")
    candidate_path = os.path.join(output_dir, "_ai_fused_candidate.jpg")
    ai_img = _ensure_min_report_size(ai_img)
    ai_img.save(candidate_path, "JPEG", quality=95)

    # 旧本地照片化兜底已停用：AI 融合失败或审查失败时直接报错。
    # final_img = _make_paper_photo(merged_path, seed=20260530 + attempt)
    log("步骤 4: 审查 AI 融合最终图片...")
    ok, report = validate_template_image(
        candidate_path,
        sub_name=sub_name,
        episode_count=episode_count,
        total_min=total_min,
    )
    if not ok:
        raise RuntimeError(f"AI 融合成本配置模板审查未通过: {report}")

    shutil.copy2(candidate_path, out_path)
    _cleanup_intermediate_files(output_dir)
    log(f"  审查通过: {report}")
    log(f"成本配置模板完成: {out_path}")
    return out_path


def _cleanup_intermediate_files(output_dir: str) -> None:
    for pattern in (
        "_merged_input.*",
        "_filled_template.*",
        "_template_candidate_*.*",
        "_ai_naturalized.*",
        "_docx_filled_input.*",
        "_ai_fused_candidate.*",
    ):
        for path in Path(output_dir).glob(pattern):
            try:
                path.unlink()
            except OSError:
                pass


def validate_template_image(
    image_path: str | Path,
    sub_name: str,
    episode_count: int,
    total_min: int,
) -> tuple[bool, str]:
    """审查最终图片是否满足上传材料基本规范。"""
    path = Path(image_path)
    if not path.exists() or path.stat().st_size < 50_000:
        return False, "图片不存在或文件过小"

    try:
        img = Image.open(path).convert("RGB")
    except Exception as exc:
        return False, f"图片无法打开: {exc}"

    width, height = img.size
    checks: list[str] = []
    if width < 1400 or height < 1900:
        checks.append(f"分辨率不足({width}x{height})")

    # 红章区域应有足够红色像素。
    red_pixels = 0
    dark_pixels = 0
    non_white_pixels = 0
    sample_step = 4
    for y in range(0, height, sample_step):
        for x in range(0, width, sample_step):
            r, g, b = img.getpixel((x, y))
            if r > 130 and g < 100 and b < 100 and r - max(g, b) > 45:
                red_pixels += 1
            if r < 95 and g < 95 and b < 95:
                dark_pixels += 1
            if min(r, g, b) < 238:
                non_white_pixels += 1

    total_samples = (width // sample_step) * (height // sample_step)
    red_ratio = red_pixels / max(1, total_samples)
    dark_ratio = dark_pixels / max(1, total_samples)
    non_white_ratio = non_white_pixels / max(1, total_samples)

    if red_ratio < 0.0012:
        checks.append(f"红章区域不足({red_ratio:.4f})")
    if dark_ratio < 0.004:
        checks.append(f"文字/表格黑色像素不足({dark_ratio:.4f})")
    if non_white_ratio < 0.05:
        checks.append(f"图片过白或内容过少({non_white_ratio:.4f})")

    if checks:
        return False, "；".join(checks)
    return True, f"通过({width}x{height}, red={red_ratio:.4f}, dark={dark_ratio:.4f})"


def _make_paper_photo(image_path: str, seed: int = 20260530) -> Image.Image:
    """本地照片化：保留全部文字和印章，只添加木纹桌面、透视、阴影和纸张质感。"""
    paper = Image.open(image_path).convert("RGB")
    target_w, target_h = 1440, 2048
    table = _make_wood_background(target_w, target_h)

    # A4 竖纸放入 2K 竖版画布。
    paper_h = 1880
    paper_w = int(paper_h * paper.width / paper.height)
    paper = paper.resize((paper_w, paper_h), Image.LANCZOS)
    paper = _add_paper_texture(paper, seed + 1)

    rng = random.Random(seed)
    angle = rng.uniform(-1.4, -0.5)
    rotated = paper.rotate(angle, expand=True, resample=Image.BICUBIC, fillcolor=(0, 0, 0))
    alpha = Image.new("L", paper.size, 255).rotate(angle, expand=True, resample=Image.BICUBIC, fillcolor=0)

    x = (target_w - rotated.width) // 2
    y = (target_h - rotated.height) // 2 + rng.randint(10, 24)

    shadow = Image.new("RGBA", table.size, (0, 0, 0, 0))
    shadow_mask = alpha.filter(ImageFilter.GaussianBlur(28))
    shadow_layer = Image.new("RGBA", rotated.size, (0, 0, 0, 82))
    shadow.paste(shadow_layer, (x + 20, y + 24), shadow_mask)

    result = table.convert("RGBA")
    result.alpha_composite(shadow)
    result.paste(rotated.convert("RGBA"), (x, y), alpha)
    return result.convert("RGB")


def _make_wood_background(width: int, height: int) -> Image.Image:
    ref_path = _resolve_photo_reference()
    if ref_path:
        ref = Image.open(ref_path).convert("RGB")
        # 只抽取参考图顶部木纹区域，避免把参考图里的白纸也当作背景带进来。
        band_h = max(24, int(ref.height * 0.035))
        band = ref.crop((0, 0, ref.width, band_h)).resize((width, band_h), Image.LANCZOS)
        bg = Image.new("RGB", (width, height))
        for y in range(0, height, band_h):
            bg.paste(band, (0, y))
        bg = bg.crop((0, 0, width, height))
        return _enhance_wood_grain(bg)

    bg = Image.new("RGB", (width, height), (204, 176, 127))
    draw = ImageDraw.Draw(bg)
    rng = random.Random(20260530)
    for x in range(width):
        wave = math.sin(x / 32.0) * 10 + math.sin(x / 103.0) * 22
        base = 184 + int(wave)
        color = (min(230, base + 28), min(205, base + 8), max(105, base - 56))
        draw.line((x, 0, x, height), fill=color)
    for _ in range(180):
        x = rng.randint(0, width)
        color = (145 + rng.randint(0, 35), 112 + rng.randint(0, 30), 66 + rng.randint(0, 22))
        draw.line((x, 0, x + rng.randint(-8, 8), height), fill=color, width=rng.choice((1, 1, 2)))
    return bg.filter(ImageFilter.GaussianBlur(1.0))


def _resolve_photo_reference() -> str | None:
    candidates = [
        os.environ.get("RECREATE_WOOD_REFERENCE", ""),
        str(SUCAI_DIR / "照片参考.png"),
        str(SUCAI_DIR / "照片参考.jpg"),
    ]
    for candidate in candidates:
        if candidate and os.path.exists(candidate):
            return candidate
    return None


def _enhance_wood_grain(bg: Image.Image) -> Image.Image:
    bg = bg.convert("RGB")
    draw = ImageDraw.Draw(bg, "RGBA")
    rng = random.Random(20260532)
    width, height = bg.size
    for _ in range(260):
        y = rng.randint(0, height)
        alpha = rng.randint(10, 28)
        color = rng.choice(((95, 58, 24, alpha), (255, 235, 190, alpha)))
        offset = rng.randint(-80, 80)
        draw.line((0, y, width, y + offset), fill=color, width=rng.choice((1, 1, 2)))
    return bg.filter(ImageFilter.GaussianBlur(0.35))


def _add_paper_texture(paper: Image.Image, seed: int = 20260531) -> Image.Image:
    paper = paper.convert("RGB")
    px = paper.load()
    rng = random.Random(seed)
    for _ in range((paper.width * paper.height) // 80):
        x = rng.randrange(paper.width)
        y = rng.randrange(paper.height)
        r, g, b = px[x, y]
        delta = rng.randint(-5, 5)
        px[x, y] = (
            max(0, min(255, r + delta)),
            max(0, min(255, g + delta)),
            max(0, min(255, b + delta)),
        )
    return paper


def _resolve_stamp_image(template_path: str, stamp_image: str | Path | None) -> str:
    if stamp_image and os.path.exists(str(stamp_image)):
        return str(stamp_image)
    template_dir = os.path.dirname(template_path)
    for name in ("模板.jpg", "模版.jpg", "图2.jpg", "模板.png", "模版.png", "图2.png"):
        candidate = os.path.join(template_dir, name)
        if os.path.exists(candidate):
            return candidate
    return str(stamp_image or os.path.join(template_dir, "模板.jpg"))


# ============================================================
# 从 docx 提取报审机构
# ============================================================

def _extract_company_from_docx(docx_path: str | Path) -> str | None:
    """从 docx 底部段落中提取报审机构名称。"""
    try:
        doc = Document(str(docx_path))
    except Exception:
        return None
    for para in doc.paragraphs:
        text = para.text.strip()
        if "报审机构" in text:
            for sep in ("：", ":"):
                if sep in text:
                    company = text.split(sep, 1)[1].strip()
                    if company:
                        return company
    return None


# ============================================================
# 步骤 1: 填写 docx
# ============================================================

def _fill_docx(template_path: str | Path, sub_name: str, episode_count: int, total_min: int) -> str:
    """在模板 docx 的表格中填入微短剧名称和集数/时长，保存到临时文件。"""
    doc = Document(str(template_path))

    # 表格第二行：[序号, 微短剧名称, 集数和总时长, 总投资额, 演员总片酬占比]
    table = doc.tables[0]
    row = table.rows[1]

    # 填写"微短剧名称"（cell[1]）
    name_cell = row.cells[1]
    name_cell.paragraphs[0].clear()
    run = name_cell.paragraphs[0].add_run(sub_name)
    run.font.name = "仿宋_GB2312"
    run.font.size = 177800  # 14pt (EMUs)

    # 填写"集数和总时长"（cell[2]）
    duration_cell = row.cells[2]
    duration_cell.paragraphs[0].clear()
    duration_text = f" {episode_count}集，共{total_min}分钟"
    run = duration_cell.paragraphs[0].add_run(duration_text)
    run.font.name = "仿宋_GB2312"
    run.font.size = 177800

    # 保存到临时文件
    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    doc.save(tmp.name)
    tmp.close()
    return tmp.name


# ============================================================
# 步骤 2: docx → jpg
# ============================================================

def _docx_to_image(docx_path: str, log_cb=None) -> Image.Image:
    """将 docx 转为 PIL Image。必须走 Word/docx2pdf，高保真失败就直接报错。"""
    pdf_path = docx_path.replace(".docx", ".pdf")

    if log_cb:
        log_cb("  正在调用本机 Word 转换 docx，最长等待 90 秒...")
    _convert_docx_to_pdf_with_timeout(docx_path, pdf_path, timeout_sec=90)

    if not os.path.exists(pdf_path) or os.path.getsize(pdf_path) <= 0:
        raise RuntimeError("docx 转 PDF 失败：未生成有效 PDF")

    img = _pdf_to_image(pdf_path)
    for p in [docx_path, pdf_path]:
        try:
            os.unlink(p)
        except OSError:
            pass
    return img

    # 旧低保真兜底已停用：不再用 PIL 简易渲染 docx。
    # return _render_docx_as_image(docx_path)


def _convert_docx_to_pdf_with_timeout(docx_path: str, pdf_path: str, timeout_sec: int = 90) -> None:
    """用 docx2pdf/Word 转 PDF；Word 卡住时超时失败，不切换兜底。"""
    result_queue: queue.Queue[tuple[bool, BaseException | None]] = queue.Queue(maxsize=1)

    def worker() -> None:
        pythoncom = None
        try:
            try:
                import pythoncom as _pythoncom
                pythoncom = _pythoncom
                pythoncom.CoInitialize()
            except Exception:
                pythoncom = None

            from docx2pdf import convert
            convert(docx_path, pdf_path)
            result_queue.put((True, None))
        except BaseException as exc:
            result_queue.put((False, exc))
        finally:
            if pythoncom is not None:
                try:
                    pythoncom.CoUninitialize()
                except Exception:
                    pass

    thread = threading.Thread(target=worker, name="docx2pdf-convert", daemon=True)
    thread.start()
    thread.join(timeout_sec)

    if thread.is_alive():
        raise RuntimeError(
            "docx 转 PDF 超时：本机 Word 可能未安装、未激活，或弹出了需要人工确认的窗口。"
            "请关闭所有 Word 窗口，确认 Microsoft Word 可正常打开后重试。"
        )

    ok, exc = result_queue.get() if not result_queue.empty() else (False, RuntimeError("docx2pdf 未返回结果"))
    if not ok:
        raise RuntimeError(f"docx 转 PDF 失败，请确认本机 Word/docx2pdf 可用: {exc}") from exc


def _render_docx_as_image(docx_path: str) -> Image.Image:
    """用 PIL 把 docx 内容渲染成图片（简单文档/表格可用）。"""
    from docx import Document as DocxDocument

    doc = DocxDocument(docx_path)

    # 画布参数
    width, margin = 2480, 120  # A4 @300dpi 左右边距
    y = 100
    line_height = 45
    line_spacer = 20
    lines: list[tuple[str, str, int]] = []  # (text, font_name, font_size_pt)

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            lines.append(("", "", 0))
            continue
        for run in para.runs:
            if run.text.strip():
                font_size = int(run.font.size / 12700) if run.font.size else 14
                lines.append((run.text.strip(), run.font.name or "仿宋_GB2312", font_size))

    # 计算高度
    total_h = y + len(lines) * (line_height + line_spacer) + 200

    img = Image.new("RGB", (width, total_h), (255, 255, 255))
    from PIL import ImageDraw, ImageFont
    draw = ImageDraw.Draw(img)

    for text, font_name, font_size in lines:
        if not text:
            y += line_spacer
            continue
        try:
            font = ImageFont.truetype(font_name, int(font_size * 2.5))
        except Exception:
            try:
                font = ImageFont.truetype("simfang.ttf", int(font_size * 2.5))
            except Exception:
                font = ImageFont.load_default()
        draw.text((margin, y), text, fill=(0, 0, 0), font=font)
        y += line_height

    # 渲染表格
    for table in doc.tables:
        y += 30
        for ri, row in enumerate(table.rows):
            col_count = len(row.cells)
            col_w = (width - 2 * margin) // col_count
            for ci, cell in enumerate(row.cells):
                x0 = margin + ci * col_w
                x1 = x0 + col_w
                y0 = y
                y1 = y + 50
                draw.rectangle([x0, y0, x1, y1], outline=(0, 0, 0), width=2)
                cell_text = cell.text.strip()
                if cell_text:
                    try:
                        cell_font = ImageFont.truetype("simfang.ttf", 35)
                    except Exception:
                        cell_font = ImageFont.load_default()
                    draw.text((x0 + 8, y0 + 8), cell_text, fill=(0, 0, 0), font=cell_font)
            y += 50

    # 清理
    try:
        os.unlink(docx_path)
    except OSError:
        pass

    return img
def _pdf_to_image(pdf_path: str) -> Image.Image:
    """使用 PyMuPDF 将 PDF 第一页转为 PIL Image。"""
    import fitz

    doc = fitz.open(pdf_path)
    page = doc[0]
    # 高分辨率渲染 (300 DPI)
    mat = fitz.Matrix(300 / 72, 300 / 72)
    pix = page.get_pixmap(matrix=mat)
    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
    doc.close()
    return img


# ============================================================
# 步骤 3: 融合图片
# ============================================================

def _merge_images(docx_img: Image.Image, stamp_image_path: str) -> Image.Image:
    """
    上半部分取 docx 生成的清晰图片，下半部分取图2.jpg 的印章区域。
    以图片高度的 65% 作为分割点。
    """
    stamp_img = Image.open(stamp_image_path)

    # 统一宽度
    target_width = max(docx_img.width, stamp_img.width)
    if docx_img.width != target_width:
        ratio = target_width / docx_img.width
        docx_img = docx_img.resize((target_width, int(docx_img.height * ratio)), Image.LANCZOS)
    if stamp_img.width != target_width:
        ratio = target_width / stamp_img.width
        stamp_img = stamp_img.resize((target_width, int(stamp_img.height * ratio)), Image.LANCZOS)

    # 分割点：docx 图片的 65% 处
    split_y = int(docx_img.height * 0.65)
    final_height = split_y + (stamp_img.height - int(stamp_img.height * 0.65))

    merged = Image.new("RGB", (target_width, final_height), (255, 255, 255))
    # 上半部分：docx 图片
    top_crop = docx_img.crop((0, 0, target_width, split_y))
    merged.paste(top_crop, (0, 0))
    # 下半部分：图2.jpg 的印章区域
    stamp_split_y = int(stamp_img.height * 0.65)
    bottom_crop = stamp_img.crop((0, stamp_split_y, target_width, stamp_img.height))
    merged.paste(bottom_crop, (0, split_y))

    return merged


def _fill_stamp_template(
    stamp_image_path: str,
    output_dir: str,
    sub_name: str,
    episode_count: int,
    total_min: int,
    log,
) -> str | None:
    """在完整模板图的表格空格中直接填写剧名和集数时长，作为稳定兜底输入。"""
    try:
        img = Image.open(stamp_image_path).convert("RGB")
    except Exception as exc:
        log(f"  模板图兜底失败: {exc}")
        return None

    draw = ImageDraw.Draw(img)
    width, height = img.size

    # 按当前模板图的表格比例定位：序号、剧名、集数时长、投资额、片酬占比。
    # 这些比例来自标准 A4 成本配置模板，能随图片尺寸等比缩放。
    name_box = _scale_box((224, 438, 396, 589), width, height)
    duration_box = _scale_box((398, 438, 596, 589), width, height)

    _clear_cell(draw, name_box)
    _clear_cell(draw, duration_box)

    name_font = _load_cn_font(max(22, int(width * 0.026)))
    duration_font = _load_cn_font(max(22, int(width * 0.025)))
    _draw_centered_text(draw, name_box, sub_name, name_font)
    _draw_centered_text(draw, duration_box, f"{episode_count}集，共{total_min}分钟", duration_font)

    out_path = os.path.join(output_dir, "_filled_template.jpg")
    img.save(out_path, "JPEG", quality=95)
    log(f"  已生成稳定模板输入: {out_path}")
    return out_path


def _scale_box(box, width: int, height: int):
    base_w, base_h = 1170, 1671
    x0, y0, x1, y1 = box
    return (
        int(x0 * width / base_w),
        int(y0 * height / base_h),
        int(x1 * width / base_w),
        int(y1 * height / base_h),
    )


def _clear_cell(draw: ImageDraw.ImageDraw, box):
    x0, y0, x1, y1 = box
    draw.rectangle((x0 + 4, y0 + 4, x1 - 4, y1 - 4), fill=(250, 250, 248))


def _load_cn_font(size: int):
    for font_path in (
        r"C:\Windows\Fonts\simfang.ttf",
        r"C:\Windows\Fonts\simsun.ttc",
        r"C:\Windows\Fonts\msyh.ttc",
    ):
        try:
            return ImageFont.truetype(font_path, size)
        except Exception:
            continue
    return ImageFont.load_default()


def _draw_centered_text(draw: ImageDraw.ImageDraw, box, text: str, font):
    x0, y0, x1, y1 = box
    max_width = max(20, x1 - x0 - 18)
    lines = _wrap_text(draw, text, font, max_width)
    line_heights = []
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        line_heights.append(bbox[3] - bbox[1])
    total_h = sum(line_heights) + max(0, len(lines) - 1) * 8
    y = y0 + ((y1 - y0) - total_h) // 2
    for line, line_h in zip(lines, line_heights):
        bbox = draw.textbbox((0, 0), line, font=font)
        text_w = bbox[2] - bbox[0]
        x = x0 + ((x1 - x0) - text_w) // 2
        draw.text((x, y), line, fill=(25, 25, 25), font=font)
        y += line_h + 8


def _wrap_text(draw: ImageDraw.ImageDraw, text: str, font, max_width: int):
    lines = []
    current = ""
    for char in text:
        test = current + char
        bbox = draw.textbbox((0, 0), test, font=font)
        if bbox[2] - bbox[0] <= max_width or not current:
            current = test
        else:
            lines.append(current)
            current = char
    if current:
        lines.append(current)
    return lines


# ============================================================
# 步骤 4: AI 生成自然照片效果
# ============================================================

def _ai_fuse_template_photo(
    docx_image_path: str,
    stamp_image_path: str,
    sub_name: str,
    episode_count: int,
    total_min: int,
    company_name: str | None = None,
    api_key: str | None = None,
    api_base_url: str | None = None,
    image_model: str | None = None,
    log_cb=None,
) -> Image.Image:
    """用 AI 把 docx 内容和印章模板融合为一张真实纸质照片。"""
    def log(msg):
        if log_cb:
            log_cb(msg)

    api_key = api_key or ""
    if not api_key:
        raise RuntimeError("未配置图片 API Key，无法生成成本配置表")

    api_base_url = _normalize_api_base_url(api_base_url)
    image_model = image_model or IMAGE_MODEL
    url = f"{api_base_url}/images/edits"

    lq = "“"
    rq = "”"
    company_instruction = (
        f"报审机构必须是：{company_name}。"
        if company_name
        else "报审机构按第二张带章模板图中的文字保留。"
    )
    prompt = (
        "请把两张输入图融合成一张真实手机拍摄照片，不要拼接痕迹，不要覆盖痕迹。\n\n"
        "输入图说明：\n"
        "1. 第一张图是已经在 docx 中改好字的成本配置比例情况报告，必须作为正文、表格和剧目信息的唯一内容来源。\n"
        "2. 第二张图是带红色圆形公章、方章、签名章和底部机构信息的模板图，底部印章、签名章、机构和日期只能来自这张图。\n\n"
        "最终图要求：\n"
        f"1. 输出是一整张完整白色 A4 打印纸，不是两张照片叠加，不是截图，不是扫描件。\n"
        f"2. 顶部标题必须为{lq}成本配置比例情况报告{rq}。\n"
        f"3. 正文必须包含{lq}我公司提审的一部“其他微短剧”，详情见下：{rq}。\n"
        "4. 表格结构和表格线必须清晰，表头为：序号、微短剧名称、集数和总时长、总投资额（需<100万）、演员总片酬占比（需<40%）。\n"
        f"5. 表格第一行必须是：1、{sub_name}、{episode_count}集，共{total_min}分钟、1万元、30%。\n"
        "6. 中下部承诺文字按第一张 docx 图保留，文字清晰、中文不乱码。\n"
        f"7. 底部必须融合第二张模板图中的红色圆形公章、红色方章/签名章、法定代表人或总编辑签名、日期位置。{company_instruction}\n"
        "8. 所有文字、表格线、印章都要像同一张纸上真实打印/盖章形成的内容，不能出现灰色擦除块、局部贴图边界、上下拼接线、重影、错位。\n"
        "9. 纸张放在浅色木纹桌面上，俯拍略带透视，纸张上边缘略远、下边缘略近，有自然阴影、轻微纸张纹理和真实手机拍摄感，画面干净，只保留纸和桌面，纸张四个角都要露出来。\n"
        "10. 不要生成新的印章，不要参考任何其他印章样式，所有红章、签名章和章中文字必须以第二张模板图为准。\n"
    )

    paths = [docx_image_path, stamp_image_path]
    for path in paths:
        if not os.path.exists(path):
            raise FileNotFoundError(f"AI 融合输入图不存在: {path}")

    files = []
    handles = []
    try:
        for path in paths:
            handle = open(path, "rb")
            handles.append(handle)
            mime_type = mimetypes.guess_type(path)[0] or "image/jpeg"
            files.append(("image", (os.path.basename(path), handle, mime_type)))
        data = {
            "model": image_model,
            "prompt": prompt,
            "size": "1024x1536",
            "n": 1,
            "response_format": "b64_json",
        }
        log(f"  请求 {url} ...")
        response = _session.post(
            url,
            headers={"Authorization": f"Bearer {api_key}"},
            files=files,
            data=data,
            timeout=300,
        )
    finally:
        for handle in handles:
            handle.close()

    if response.status_code != 200:
        raise RuntimeError(f"AI 融合请求失败 HTTP {response.status_code}: {response.text[:500]}")

    img = _decode_image_response(response)
    log("  AI 融合完成")
    return img


def _decode_image_response(response: requests.Response) -> Image.Image:
    try:
        result = response.json()
    except Exception as exc:
        raise RuntimeError(f"AI 返回不是 JSON: {exc}") from exc

    images = result.get("data", [])
    if not images:
        raise RuntimeError(f"AI 返回缺少 data 图片: {str(result)[:500]}")

    first = images[0]
    if "b64_json" in first and first["b64_json"]:
        img_bytes = base64.b64decode(first["b64_json"])
        return Image.open(BytesIO(img_bytes)).convert("RGB")

    if "url" in first and first["url"]:
        resp = _session.get(first["url"], timeout=120)
        if resp.status_code != 200:
            raise RuntimeError(f"下载 AI 图片失败 HTTP {resp.status_code}: {resp.text[:200]}")
        return Image.open(BytesIO(resp.content)).convert("RGB")

    raise RuntimeError(f"AI 返回图片格式不支持: {list(first.keys())}")


def _ensure_min_report_size(img: Image.Image) -> Image.Image:
    """AI 结果只做尺寸放大，不改内容；确保后续上传审查有足够分辨率。"""
    img = img.convert("RGB")
    min_w, min_h = 1440, 2048
    if img.width >= min_w and img.height >= min_h:
        return img

    scale = max(min_w / img.width, min_h / img.height)
    new_size = (math.ceil(img.width * scale), math.ceil(img.height * scale))
    return img.resize(new_size, Image.LANCZOS)


def _ai_naturalize(
    image_path: str,
    sub_name: str,
    episode_count: int,
    total_min: int,
    company_name: str | None = None,
    api_key: str | None = None,
    api_base_url: str | None = None,
    image_model: str | None = None,
    log_cb=None,
) -> Image.Image | None:
    """
    使用 images/edits API 生成真实照片效果的完整纸质文件。
    """
    def log(msg):
        if log_cb:
            log_cb(msg)

    api_key = api_key or ""
    if not api_key:
        log("  未配置图片 API Key，跳过 AI 自然化")
        return None

    api_base_url = _normalize_api_base_url(api_base_url)
    image_model = image_model or IMAGE_MODEL
    url = f"{api_base_url}/images/edits"

    # 报审机构：优先使用从 docx 提取的公司名，未提取到则让 AI 照着输入图自行保留。
    lq = "“"
    rq = "”"
    company_instruction = (
        f"底部保留报审机构：{company_name}，保留法定代表人或总编辑签名、红色方章、红色圆形公司公章和{lq}2026 年  月  日{rq}。"
        if company_name
        else f"底部保留报审机构、法定代表人或总编辑签名、红色方章、红色圆形公司公章和{lq}2026 年  月  日{rq}，全部照着输入图原样保留。"
    )

    prompt = (
        "请将输入图中的中文报告处理成一张真实纸质文件照片。\n\n"
        "文件必须是一张完整白色 A4 纸，放在浅色木纹桌面上，从上往下略带透视角度，"
        "像手机真实拍摄。纸张有轻微纹理、自然阴影和边缘透视感。\n\n"
        "内容要求：\n"
        f"1. 顶部标题必须是{lq}成本配置比例情况报告{rq}。\n"
        f"2. 正文保留{lq}我公司提审的一部'其他微短剧'，详情见下：{rq}。\n"
        "3. 保留表格结构和表格线，表头为：序号、微短剧名称、集数和总时长、总投资额（需<100万）、演员总片酬占比（需<40%）。\n"
        f"4. 表格第一行内容为：1、{sub_name}、{episode_count}集，共{total_min}分钟、1万元、30%。\n"
        "5. 中下部保留承诺文字：以上微短剧制作完成的总投资额和演员总片酬占比均符合《关于电视剧网络剧制作成本配置比例的意见》要求。"
        "我公司郑重承诺拥有上述一部剧目的播出版权，并承诺上述内容真实有效。如因材料真实性、演员片酬占比等问题产生的不良后果，我公司愿承担相关责任。\n"
        f"6. {company_instruction}\n"
        "7. 去掉任何绿色背景、编辑器按钮、拖拽控件、边框辅助线、AI按钮等界面元素。\n\n"
        "视觉要求：文字清晰可读，中文不要乱码；黑色打印字体自然；红章真实；表格线清晰；"
        "不要电子截图效果，不要扫描件效果。最终效果是一张真实打印出来的成本配置比例情况报告照片。"
    )

    try:
        with open(image_path, "rb") as f:
            files = {
                "image": (os.path.basename(image_path), f, "image/jpeg"),
            }
            data = {
                "model": image_model,
                "prompt": prompt,
                "size": "1024x1024",
                "n": 1,
            }
            log(f"  请求 {url} ...")
            response = _session.post(
                url,
                headers={"Authorization": f"Bearer {api_key}"},
                files=files,
                data=data,
                timeout=300,
            )

        if response.status_code == 200:
            result = response.json()
            images = result.get("data", [])
            if images and "b64_json" in images[0]:
                img_bytes = base64.b64decode(images[0]["b64_json"])
                log("  AI 自然化完成")
                return Image.open(BytesIO(img_bytes))
            log(f"  API 返回异常: {list(images[0].keys()) if images else 'empty'}")
        else:
            log(f"  AI 请求失败 HTTP {response.status_code}: {response.text[:200]}")
    except Exception as exc:
        log(f"  AI 生图异常: {exc}")

    return None
